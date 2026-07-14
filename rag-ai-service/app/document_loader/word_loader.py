from __future__ import annotations

from io import BytesIO

from docx import Document as DocxDocument

from app.document_loader.base import ParsedDocument, ParsedPage


def load_word(file_bytes: bytes) -> ParsedDocument:
    """从 DOCX 段落和简单表格中提取文本。

    如果没有排版引擎，DOCX 很难得到稳定渲染页码，
    所以这里把整份文档视为一个逻辑段落，page_no=None。
    """
    document = DocxDocument(BytesIO(file_bytes))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            # 把表格行压平成可读文本，避免制度/规格表格在检索时被静默丢弃。
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                parts.append(" | ".join(values))

    return ParsedDocument(pages=[ParsedPage(page_no=None, text="\n".join(parts))])
