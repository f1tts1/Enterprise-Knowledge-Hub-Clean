from __future__ import annotations

import os
import unittest
from io import BytesIO


os.environ.setdefault("MINIO_ACCESS_KEY", "unit-test-access-key")
os.environ.setdefault("MINIO_SECRET_KEY", "unit-test-secret-key")

from app.document_loader.base import ParsedDocument, ParsedPage  # noqa: E402
from app.document_loader.factory import (  # noqa: E402
    UnsupportedDocumentTypeError,
    load_document,
)
from app.splitter.recursive_splitter import RecursiveTextSplitter  # noqa: E402


class DocumentLoaderTest(unittest.TestCase):
    def test_text_and_markdown_loaders_preserve_decoded_content(self) -> None:
        utf8_document = load_document("guide.TXT", "知识库内容".encode("utf-8"))
        gb18030_document = load_document("legacy.txt", "旧版制度".encode("gb18030"))
        markdown_document = load_document("README.markdown", b"# Heading\n\nBody")

        self.assertEqual("知识库内容", utf8_document.text)
        self.assertEqual("旧版制度", gb18030_document.text)
        self.assertEqual("# Heading\n\nBody", markdown_document.text)
        self.assertEqual(1, utf8_document.pages[0].page_no)

    def test_pdf_loader_preserves_page_numbers(self) -> None:
        import fitz

        source = fitz.open()
        first_page = source.new_page()
        first_page.insert_text((72, 72), "first page evidence")
        second_page = source.new_page()
        second_page.insert_text((72, 72), "second page evidence")
        payload = source.tobytes()
        source.close()

        parsed = load_document("evidence.pdf", payload)

        self.assertEqual(2, parsed.page_count)
        self.assertEqual([1, 2], [page.page_no for page in parsed.pages])
        self.assertIn("first page evidence", parsed.pages[0].text)
        self.assertIn("second page evidence", parsed.pages[1].text)

    def test_docx_loader_extracts_paragraphs_and_table_rows(self) -> None:
        from docx import Document

        source = Document()
        source.add_paragraph("合同归档规则")
        table = source.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "期限"
        table.cell(0, 1).text = "五个工作日"
        buffer = BytesIO()
        source.save(buffer)

        parsed = load_document("policy.docx", buffer.getvalue())

        self.assertEqual(1, parsed.page_count)
        self.assertIsNone(parsed.pages[0].page_no)
        self.assertIn("合同归档规则", parsed.text)
        self.assertIn("期限 | 五个工作日", parsed.text)

    def test_unsupported_extension_is_rejected(self) -> None:
        with self.assertRaises(UnsupportedDocumentTypeError):
            load_document("archive.zip", b"not-a-document")


class RecursiveTextSplitterTest(unittest.TestCase):
    def test_invalid_chunk_configuration_is_rejected(self) -> None:
        invalid_configs = (
            {"chunk_size": 0, "chunk_overlap": 0},
            {"chunk_size": 20, "chunk_overlap": -1},
            {"chunk_size": 20, "chunk_overlap": 20},
        )

        for config in invalid_configs:
            with self.subTest(config=config), self.assertRaises(ValueError):
                RecursiveTextSplitter(**config)

    def test_splitter_keeps_chunk_bounds_and_page_metadata(self) -> None:
        document = ParsedDocument(
            pages=[
                ParsedPage(page_no=1, text="第一条规则。第二条规则。第三条规则。第四条规则。"),
                ParsedPage(page_no=2, text="   "),
                ParsedPage(page_no=3, text="alpha beta gamma delta epsilon zeta eta theta"),
            ]
        )
        splitter = RecursiveTextSplitter(chunk_size=18, chunk_overlap=4)

        chunks = splitter.split_document(document)

        self.assertGreater(len(chunks), 2)
        self.assertEqual(list(range(len(chunks))), [chunk.chunk_index for chunk in chunks])
        self.assertEqual({1, 3}, {chunk.page_no for chunk in chunks})
        for chunk in chunks:
            with self.subTest(chunk=chunk):
                self.assertTrue(chunk.text)
                self.assertLessEqual(len(chunk.text), 18)
                self.assertGreaterEqual(chunk.char_start, 0)
                self.assertEqual(len(chunk.text), chunk.char_end - chunk.char_start)


if __name__ == "__main__":
    unittest.main()
